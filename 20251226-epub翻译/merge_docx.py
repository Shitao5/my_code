#!/usr/bin/env python3
import argparse
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxcompose.composer import Composer


def set_songti_run(run):
    run.font.name = "宋体"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "宋体")


def set_songti_style(style):
    if not hasattr(style, "font"):
        return
    style.font.name = "宋体"
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "宋体")


def iter_paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def has_drawing(paragraph):
    return bool(paragraph._element.xpath(".//w:drawing | .//w:pict"))


def run_has_drawing(run):
    return bool(run._element.xpath(".//w:drawing | .//w:pict"))


def last_text_run(paragraph):
    for run in reversed(paragraph.runs):
        if run_has_drawing(run):
            continue
        if run.text:
            return run
    return None


def normalize_paragraphs(doc):
    paragraphs = list(iter_paragraphs(doc))
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if run_has_drawing(run):
                continue
            if "</translation>" in run.text:
                run.text = run.text.replace("</translation>", "")
            run.text = run.text.replace("\r\n", "\n").replace("\r", "\n")
            run.text = re.sub(r"\n{2,}", "\n", run.text)

    for idx, paragraph in enumerate(paragraphs[:-1]):
        run = last_text_run(paragraph)
        if run is None:
            continue
        run.text = run.text.rstrip("\n")

    prev_empty = False
    for paragraph in list(iter_paragraphs(doc)):
        text = "".join(run.text for run in paragraph.runs).strip()
        is_empty = text == "" and not has_drawing(paragraph)
        if is_empty and prev_empty:
            paragraph._element.getparent().remove(paragraph._element)
            continue
        prev_empty = is_empty


def normalize_images(doc, width_cm=14):
    for paragraph in iter_paragraphs(doc):
        if has_drawing(paragraph):
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for shape in doc.inline_shapes:
        shape.width = Cm(width_cm)


def merge_docx(input_dir, output_path):
    files = sorted(Path(input_dir).glob("*.docx"))
    if not files:
        raise SystemExit(f"No .docx files found in {input_dir}")

    base_doc = Document(files[0])
    composer = Composer(base_doc)
    for file_path in files[1:]:
        composer.append(Document(file_path))

    doc = composer.doc
    normalize_paragraphs(doc)
    normalize_images(doc, width_cm=14)
    for style in doc.styles:
        set_songti_style(style)
    for paragraph in iter_paragraphs(doc):
        for run in paragraph.runs:
            set_songti_run(run)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(
        description="Merge .docx files and normalize line breaks."
    )
    parser.add_argument(
        "-i",
        "--input-dir",
        default="out",
        help="Directory containing .docx files.",
    )
    parser.add_argument(
        "-o",
        "--output",
        default="merged.docx",
        help="Output .docx path.",
    )
    args = parser.parse_args()
    merge_docx(args.input_dir, args.output)


if __name__ == "__main__":
    main()
