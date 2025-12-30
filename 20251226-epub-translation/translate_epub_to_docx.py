#!/usr/bin/env python3
import argparse
import json
import os
import posixpath
import re
import tempfile
import time
from typing import Dict, List, Optional

import requests
from tqdm import tqdm
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from ebooklib import epub, ITEM_DOCUMENT, ITEM_IMAGE


SENTENCE_SPLIT_RE = re.compile(r"(?<=[\.\!\?\u3002\uff01\uff1f])\s+")
TRANSLATION_TAG_RE = re.compile(r"<translation>(.*?)</translation>", re.IGNORECASE | re.DOTALL)
INFERENCE_PREFIXES = (
    "思考",
    "推理",
    "分析",
    "reasoning",
    "analysis",
    "thought",
    "original",
    "source",
    "translation",
)


class OllamaTranslator:
    def __init__(self, model: str, url: str, max_chars: int, retry: int) -> None:
        self.model = model
        self.url = url.rstrip("/")
        self.max_chars = max_chars
        self.retry = retry
        self._cache: Dict[str, str] = {}

    def translate_text(self, text: str) -> str:
        if not text.strip():
            return text
        cached = self._cache.get(text)
        if cached is not None:
            return cached
        if len(text) <= self.max_chars:
            translated = self._translate_once(text)
        else:
            chunks = split_text(text, self.max_chars)
            translated = "".join(self._translate_once(chunk) for chunk in chunks)
        self._cache[text] = translated
        return translated

    def _translate_once(self, text: str) -> str:
        instruction = (
            "Translate the following English text to Simplified Chinese only. "
            "Do not explain, show reasoning, or repeat the English text. "
            "Wrap the final translation inside <translation> and </translation> tags."
        )
        last_err = None
        for _ in range(self.retry + 1):
            try:
                payload = {
                    "model": self.model,
                    "prompt": f"{instruction}\n\n{text}",
                    "stream": False,
                }
                response = requests.post(
                    f"{self.url}/api/generate",
                    json=payload,
                    timeout=120,
                )
                if response.status_code == 404:
                    chat_payload = {
                        "model": self.model,
                        "messages": [
                            {"role": "system", "content": instruction},
                            {"role": "user", "content": text},
                        ],
                        "stream": False,
                    }
                    response = requests.post(
                        f"{self.url}/api/chat",
                        json=chat_payload,
                        timeout=120,
                    )
                response.raise_for_status()
                data = response.json()
                raw_text = (
                    data.get("response")
                    or data.get("message", {}).get("content")
                    or ""
                ).strip()
                cleaned = extract_translation_only(raw_text)
                return cleaned or text
            except Exception as err:  # noqa: BLE001 - keep retry simple
                last_err = err
                time.sleep(1)
        raise RuntimeError(f"Ollama translation failed: {last_err}")


def extract_translation_only(raw: str) -> str:
    if not raw:
        return ""
    match = TRANSLATION_TAG_RE.search(raw)
    if match:
        candidate = match.group(1).strip()
        if candidate:
            return candidate
    cleaned_lines: List[str] = []
    for line in raw.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        lowered = stripped.lower()
        if any(lowered.startswith(prefix) for prefix in INFERENCE_PREFIXES):
            continue
        cleaned_lines.append(stripped)
    cleaned = "\n".join(cleaned_lines).strip()
    return cleaned


def split_text(text: str, max_chars: int) -> List[str]:
    parts = SENTENCE_SPLIT_RE.split(text)
    chunks: List[str] = []
    buf = ""
    for part in parts:
        if not part:
            continue
        if len(buf) + len(part) + 1 <= max_chars:
            buf = f"{buf} {part}".strip()
        else:
            if buf:
                chunks.append(buf)
            buf = part
    if buf:
        chunks.append(buf)
    return chunks


def normalize_epub_path(path: str) -> str:
    return posixpath.normpath(path).lstrip("./")


def build_image_map(book: epub.EpubBook, temp_dir: str) -> Dict[str, str]:
    mapping: Dict[str, str] = {}
    for item in book.get_items_of_type(ITEM_IMAGE):
        name = normalize_epub_path(item.get_name())
        safe_name = name.replace("/", "__")
        out_path = os.path.join(temp_dir, safe_name)
        with open(out_path, "wb") as handle:
            handle.write(item.get_content())
        mapping[name] = out_path
    return mapping


def resolve_image_src(src: str, doc_path: str) -> Optional[str]:
    if not src or src.startswith("data:"):
        return None
    base = posixpath.dirname(doc_path)
    resolved = normalize_epub_path(posixpath.join(base, src))
    return resolved


def merge_text_segments(segments: List[Dict[str, str]]) -> List[Dict[str, str]]:
    merged: List[Dict[str, str]] = []
    for seg in segments:
        if seg["type"] == "text":
            if merged and merged[-1]["type"] == "text":
                merged[-1]["value"] += seg["value"]
            else:
                merged.append(seg)
        else:
            merged.append(seg)
    return merged


def extract_inline_segments(node: Tag) -> List[Dict[str, str]]:
    segments: List[Dict[str, str]] = []

    def walk(child):
        if isinstance(child, NavigableString):
            segments.append({"type": "text", "value": str(child)})
            return
        if not isinstance(child, Tag):
            return
        if child.name == "img":
            segments.append({"type": "img", "src": child.get("src", "")})
            return
        if child.name == "br":
            segments.append({"type": "text", "value": "\n"})
            return
        for grand in child.children:
            walk(grand)

    walk(node)
    return merge_text_segments(segments)


def add_translated_text(paragraph, text: str, translator: OllamaTranslator) -> None:
    if not text:
        return
    if not text.strip():
        paragraph.add_run(text)
        return
    translated = translator.translate_text(text)
    if "\n" in translated:
        for idx, part in enumerate(translated.split("\n")):
            paragraph.add_run(part)
            if idx < len(translated.split("\n")) - 1:
                paragraph.add_run().add_break()
    else:
        paragraph.add_run(translated)


def process_table(tag: Tag, doc: Document, translator: OllamaTranslator) -> None:
    for row in tag.find_all("tr"):
        cells = []
        for cell in row.find_all(["td", "th"]):
            cell_text = cell.get_text(" ", strip=True)
            if cell_text:
                cells.append(translator.translate_text(cell_text))
        if cells:
            doc.add_paragraph("\t".join(cells))


def process_block(
    tag: Tag,
    doc: Document,
    translator: OllamaTranslator,
    image_map: Dict[str, str],
    doc_path: str,
) -> None:
    if tag.name == "table":
        process_table(tag, doc, translator)
        return

    paragraph = doc.add_paragraph()
    if tag.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        paragraph.style = f"Heading {tag.name[1:]}"
    elif tag.name == "li":
        paragraph.style = "List Bullet"

    segments = extract_inline_segments(tag)
    for seg in segments:
        if seg["type"] == "text":
            add_translated_text(paragraph, seg["value"], translator)
            continue
        if seg["type"] == "img":
            src = resolve_image_src(seg["src"], doc_path)
            if not src:
                continue
            img_path = image_map.get(src)
            if not img_path:
                continue
            paragraph.add_run().add_picture(img_path)


def process_document_item(
    item: epub.EpubItem,
    doc: Document,
    translator: OllamaTranslator,
    image_map: Dict[str, str],
) -> None:
    content = item.get_content()
    soup = BeautifulSoup(content, "lxml-xml")
    body = soup.body
    if not body:
        return

    block_tags = {
        "p",
        "div",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "li",
        "blockquote",
        "table",
    }

    for child in body.children:
        if isinstance(child, NavigableString):
            if child.strip():
                paragraph = doc.add_paragraph()
                add_translated_text(paragraph, str(child), translator)
            continue
        if not isinstance(child, Tag):
            continue
        if child.name in block_tags:
            process_block(child, doc, translator, image_map, item.get_name())
        else:
            paragraph = doc.add_paragraph()
            segments = extract_inline_segments(child)
            for seg in segments:
                if seg["type"] == "text":
                    add_translated_text(paragraph, seg["value"], translator)
                elif seg["type"] == "img":
                    src = resolve_image_src(seg["src"], item.get_name())
                    if not src:
                        continue
                    img_path = image_map.get(src)
                    if img_path:
                        paragraph.add_run().add_picture(img_path)


def translate_epub_to_docx(
    input_path: str,
    output_dir: str,
    model: str,
    url: str,
    max_chars: int,
    retry: int,
) -> None:
    book = epub.read_epub(input_path)
    translator = OllamaTranslator(model=model, url=url, max_chars=max_chars, retry=retry)
    os.makedirs(output_dir, exist_ok=True)
    with tempfile.TemporaryDirectory() as temp_dir:
        image_map = build_image_map(book, temp_dir)
        doc_items = list(book.get_items_of_type(ITEM_DOCUMENT))
        total = len(doc_items)
        for idx, item in enumerate(
            tqdm(doc_items, desc="Translating", unit="section"), start=1
        ):
            section_doc = Document()
            process_document_item(item, section_doc, translator, image_map)
            section_path = os.path.join(
                output_dir, f"section_{idx:02d}_of_{total:02d}.docx"
            )
            section_doc.save(section_path)
            section_text_path = os.path.join(
                output_dir, f"section_{idx:02d}_of_{total:02d}.txt"
            )
            with open(section_text_path, "w", encoding="utf-8") as handle:
                handle.write(document_to_plain_text(section_doc))


def document_to_plain_text(doc: Document) -> str:
    lines: List[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        lines.append(text)
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Translate an English EPUB to Chinese DOCX using a local Ollama model."
    )
    parser.add_argument("--input", required=True, help="Path to input .epub file.")
    parser.add_argument(
        "--output-dir",
        required=True,
        help="Directory to write section_XX.docx files.",
    )
    parser.add_argument("--model", default="qwen3:8b", help="Ollama model name.")
    parser.add_argument(
        "--url",
        default="http://localhost:11434",
        help="Ollama base URL.",
    )
    parser.add_argument(
        "--max-chars",
        type=int,
        default=1500,
        help="Max characters per translation chunk.",
    )
    parser.add_argument(
        "--retry",
        type=int,
        default=2,
        help="Retry count for Ollama calls.",
    )
    args = parser.parse_args()

    translate_epub_to_docx(
        input_path=args.input,
        output_dir=args.output_dir,
        model=args.model,
        url=args.url,
        max_chars=args.max_chars,
        retry=args.retry,
    )


if __name__ == "__main__":
    main()
